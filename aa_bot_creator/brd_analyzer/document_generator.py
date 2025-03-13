#!/usr/bin/env python3
"""
Document Generator Module for AA Bot Creator.

This module provides functionality to generate various types of documents
based on requirements extracted from BRDs.
"""

import os
import sys
import logging
import json
import re
from datetime import datetime
import matplotlib.pyplot as plt
import networkx as nx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Try to import AI modules, but continue if not available
try:
    from langchain.llms import OpenAI
    from langchain.prompts import PromptTemplate
    AI_AVAILABLE = True
except ImportError:
    AI_AVAILABLE = False

class DocumentGenerator:
    """
    A class for generating documentation from requirements.
    
    This class provides methods to generate various types of documents,
    including Solution Design Documents (SDD) and User Story Documents (USD).
    It also includes methods for generating flow diagrams.
    """
    
    def __init__(self, requirements, output_dir=None):
        """
        Initialize the DocumentGenerator with requirements and output directory.
        
        Args:
            requirements (dict): The requirements extracted from the BRD.
            output_dir (str, optional): The directory to save generated documents.
                Defaults to current directory if not specified.
        """
        self.requirements = requirements
        self.output_dir = output_dir or os.getcwd()
        
        # Create output directory if it doesn't exist
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
        
        # Set up logging
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.INFO)
        
        # Set up OpenAI if available
        self.ai_model = None
        if AI_AVAILABLE:
            try:
                openai_api_key = os.environ.get("OPENAI_API_KEY")
                if openai_api_key:
                    self.ai_model = OpenAI(api_key=openai_api_key)
                    self.logger.info("AI model initialized successfully.")
                else:
                    self.logger.warning("OPENAI_API_KEY not found in environment variables.")
            except Exception as e:
                self.logger.error(f"Error initializing AI model: {str(e)}")
    
    def generate_solution_design_document(self):
        """
        Generate a Solution Design Document (SDD) based on the requirements.
        
        Returns:
            bool: True if the document was generated successfully, False otherwise.
        """
        try:
            self.logger.info("Generating Solution Design Document...")
            
            # Create a new Document
            doc = Document()
            
            # Add document title
            title = doc.add_heading('Solution Design Document', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add document information section
            doc.add_heading('Document Information', level=1)
            doc.add_paragraph(f"Project: {self.requirements.get('project_info', {}).get('name', 'N/A')}")
            doc.add_paragraph(f"Version: {self.requirements.get('project_info', {}).get('version', '1.0.0')}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
            doc.add_paragraph(f"Author: AA Bot Creator")
            
            # Add introduction section
            doc.add_heading('1. Introduction', level=1)
            doc.add_paragraph(self.requirements.get('project_info', {}).get('description', 'No description available.'))
            
            # Add purpose section
            doc.add_heading('2. Purpose', level=1)
            doc.add_paragraph('This document provides a detailed technical design for the automation solution.')
            
            # Add scope section
            doc.add_heading('3. Scope', level=1)
            doc.add_paragraph('This document covers the design of all components required for the automation solution.')
            
            # Add system overview section
            doc.add_heading('4. System Overview', level=1)
            system_overview = self._generate_system_overview()
            doc.add_paragraph(system_overview)
            
            # Add architecture section
            doc.add_heading('5. Architecture', level=1)
            doc.add_paragraph('The solution follows a modular architecture with the following components:')
            
            # Add components section
            doc.add_heading('6. Components', level=1)
            for req_type in ['data_extraction', 'data_processing', 'system_integration', 'process_automation']:
                reqs = self._get_requirements_by_type(req_type)
                if reqs:
                    doc.add_heading(f'6.{req_type.replace("_", " ").title()} Components', level=2)
                    for req in reqs:
                        doc.add_paragraph(f"{req['id']}: {req['description']}", style='List Bullet')
                        doc.add_paragraph(f"Details: {req['details']}")
            
            # Add requirements section
            doc.add_heading('7. Requirements', level=1)
            for i, req in enumerate(self.requirements.get('requirements', [])):
                doc.add_heading(f"7.{i+1}. {req['id']}: {req['description']}", level=2)
                doc.add_paragraph(f"Type: {req['type']}")
                doc.add_paragraph(f"Priority: {req['priority']}")
                doc.add_paragraph(f"Details: {req['details']}")
                if req['dependencies']:
                    doc.add_paragraph(f"Dependencies: {', '.join(req['dependencies'])}")
            
            # Add technical design section
            doc.add_heading('8. Technical Design', level=1)
            technical_design = self._generate_technical_design()
            doc.add_paragraph(technical_design)
            
            # Add implementation details section
            doc.add_heading('9. Implementation Details', level=1)
            implementation_details = self._generate_implementation_details()
            doc.add_paragraph(implementation_details)
            
            # Add testing section
            doc.add_heading('10. Testing', level=1)
            doc.add_paragraph('The solution will be tested using the following approach:')
            doc.add_paragraph('- Unit Testing: Testing individual components', style='List Bullet')
            doc.add_paragraph('- Integration Testing: Testing component interactions', style='List Bullet')
            doc.add_paragraph('- System Testing: Testing the entire solution', style='List Bullet')
            doc.add_paragraph('- User Acceptance Testing: Validation by business users', style='List Bullet')
            
            # Add deployment section
            doc.add_heading('11. Deployment', level=1)
            doc.add_paragraph('The solution will be deployed to the Automation Anywhere Control Room.')
            
            # Add maintenance section
            doc.add_heading('12. Maintenance', level=1)
            doc.add_paragraph('The solution will be maintained according to the following schedule:')
            doc.add_paragraph('- Regular monitoring of bot execution logs', style='List Bullet')
            doc.add_paragraph('- Weekly review of exceptions and failures', style='List Bullet')
            doc.add_paragraph('- Monthly performance optimization', style='List Bullet')
            
            # Add conclusion section
            doc.add_heading('13. Conclusion', level=1)
            doc.add_paragraph('This document provides a comprehensive design for the automation solution.')
            
            # Save the document
            output_path = os.path.join(self.output_dir, 'Solution_Design_Document.docx')
            doc.save(output_path)
            
            self.logger.info(f"Solution Design Document generated successfully: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error generating Solution Design Document: {str(e)}")
            return False
    
    def generate_user_story_document(self):
        """
        Generate a User Story Document (USD) based on the requirements.
        
        Returns:
            bool: True if the document was generated successfully, False otherwise.
        """
        try:
            self.logger.info("Generating User Story Document...")
            
            # Create a new Document
            doc = Document()
            
            # Add document title
            title = doc.add_heading('User Story Document', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add document information section
            doc.add_heading('Document Information', level=1)
            doc.add_paragraph(f"Project: {self.requirements.get('project_info', {}).get('name', 'N/A')}")
            doc.add_paragraph(f"Version: {self.requirements.get('project_info', {}).get('version', '1.0.0')}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
            doc.add_paragraph(f"Author: AA Bot Creator")
            
            # Add introduction section
            doc.add_heading('1. Introduction', level=1)
            doc.add_paragraph(self.requirements.get('project_info', {}).get('description', 'No description available.'))
            
            # Add purpose section
            doc.add_heading('2. Purpose', level=1)
            doc.add_paragraph('This document outlines the user stories for the automation solution.')
            
            # Add scope section
            doc.add_heading('3. Scope', level=1)
            doc.add_paragraph('This document covers all user stories related to the automation solution.')
            
            # Add user stories section
            doc.add_heading('4. User Stories', level=1)
            
            # Generate user stories for each requirement
            user_stories = self._generate_user_stories()
            
            for i, (req_id, story) in enumerate(user_stories.items()):
                doc.add_heading(f"4.{i+1}. User Story for {req_id}", level=2)
                doc.add_paragraph(story)
            
            # Add conclusion section
            doc.add_heading('5. Conclusion', level=1)
            doc.add_paragraph('This document provides a comprehensive set of user stories for the automation solution.')
            
            # Save the document
            output_path = os.path.join(self.output_dir, 'User_Story_Document.docx')
            doc.save(output_path)
            
            self.logger.info(f"User Story Document generated successfully: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error generating User Story Document: {str(e)}")
            return False
    
    def generate_flow_diagram(self):
        """
        Generate a flow diagram based on the requirements.
        
        Returns:
            bool: True if the diagram was generated successfully, False otherwise.
        """
        try:
            self.logger.info("Generating Flow Diagram...")
            
            # Create a directed graph
            G = nx.DiGraph()
            
            # Add nodes for each requirement
            for req in self.requirements.get('requirements', []):
                G.add_node(req['id'], label=f"{req['id']}\n{req['description'][:20]}...")
                
                # Add edges for dependencies
                for dep in req.get('dependencies', []):
                    G.add_edge(dep, req['id'])
            
            # Create the plot
            plt.figure(figsize=(12, 8))
            pos = nx.spring_layout(G)
            nx.draw_networkx(G, pos, with_labels=True, node_color='lightblue', 
                             node_size=2000, font_size=10, arrows=True)
            
            # Save the diagram
            output_path = os.path.join(self.output_dir, 'Flow_Diagram.png')
            plt.savefig(output_path)
            plt.close()
            
            self.logger.info(f"Flow Diagram generated successfully: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error generating Flow Diagram: {str(e)}")
            return False
    
    def _get_requirements_by_type(self, req_type):
        """
        Get requirements of a specific type.
        
        Args:
            req_type (str): The type of requirements to filter.
            
        Returns:
            list: A list of requirements of the specified type.
        """
        return [req for req in self.requirements.get('requirements', []) if req['type'] == req_type]
    
    def _generate_system_overview(self):
        """
        Generate a system overview based on the requirements.
        
        Returns:
            str: A system overview description.
        """
        if self.ai_model:
            try:
                prompt = PromptTemplate(
                    input_variables=["requirements"],
                    template="Generate a system overview based on these requirements:\n{requirements}\n\nSystem Overview:"
                )
                
                # Convert requirements to string for the prompt
                req_str = json.dumps(self.requirements, indent=2)
                
                # Generate system overview using AI
                return self.ai_model(prompt.format(requirements=req_str))
            except Exception as e:
                self.logger.error(f"Error generating system overview with AI: {str(e)}")
        
        # Fallback to template-based generation
        project_name = self.requirements.get('project_info', {}).get('name', 'the automation solution')
        project_desc = self.requirements.get('project_info', {}).get('description', 'automate business processes')
        
        return f"""
        {project_name} is designed to {project_desc}. The system consists of multiple components 
        that work together to achieve the automation goals. It includes data extraction capabilities, 
        data processing logic, system integration points, and process automation workflows.
        
        The solution is built on the Automation Anywhere platform, leveraging its robust RPA 
        capabilities to automate repetitive tasks and integrate with existing systems.
        """
    
    def _generate_technical_design(self):
        """
        Generate a technical design based on the requirements.
        
        Returns:
            str: A technical design description.
        """
        if self.ai_model:
            try:
                prompt = PromptTemplate(
                    input_variables=["requirements"],
                    template="Generate a technical design based on these requirements:\n{requirements}\n\nTechnical Design:"
                )
                
                # Convert requirements to string for the prompt
                req_str = json.dumps(self.requirements, indent=2)
                
                # Generate technical design using AI
                return self.ai_model(prompt.format(requirements=req_str))
            except Exception as e:
                self.logger.error(f"Error generating technical design with AI: {str(e)}")
        
        # Fallback to template-based generation
        return """
        The technical design follows a modular approach with the following components:
        
        1. Data Extraction Layer:
           - Responsible for extracting data from various sources
           - Uses OCR and pattern matching techniques
           - Handles different file formats and structures
        
        2. Data Processing Layer:
           - Validates and transforms extracted data
           - Applies business rules and logic
           - Handles exceptions and edge cases
        
        3. System Integration Layer:
           - Connects with external systems via APIs
           - Handles authentication and data exchange
           - Ensures data consistency across systems
        
        4. Process Automation Layer:
           - Orchestrates the end-to-end process
           - Manages workflow and dependencies
           - Provides logging and error handling
        
        The solution uses Automation Anywhere's built-in capabilities for task automation,
        along with custom scripts for complex logic.
        """
    
    def _generate_implementation_details(self):
        """
        Generate implementation details based on the requirements.
        
        Returns:
            str: Implementation details description.
        """
        if self.ai_model:
            try:
                prompt = PromptTemplate(
                    input_variables=["requirements"],
                    template="Generate implementation details based on these requirements:\n{requirements}\n\nImplementation Details:"
                )
                
                # Convert requirements to string for the prompt
                req_str = json.dumps(self.requirements, indent=2)
                
                # Generate implementation details using AI
                return self.ai_model(prompt.format(requirements=req_str))
            except Exception as e:
                self.logger.error(f"Error generating implementation details with AI: {str(e)}")
        
        # Fallback to template-based generation
        return """
        The implementation will use the following Automation Anywhere features:
        
        1. For Data Extraction:
           - Document OCR package for text extraction
           - PDF integration package for PDF processing
           - Email automation for handling email attachments
        
        2. For Data Processing:
           - String operation packages for data transformation
           - Excel advanced packages for data manipulation
           - Error handling framework for exception management
        
        3. For System Integration:
           - REST Web Service package for API integration
           - Database packages for database operations
           - Credential Vault for secure authentication
        
        4. For Process Automation:
           - Task Bot for main process flow
           - MetaBots for reusable components
           - Workload Management for scaling
        
        The solution will be implemented using a combination of pre-built actions and
        custom scripts where necessary.
        """
    
    def _generate_user_stories(self):
        """
        Generate user stories based on the requirements.
        
        Returns:
            dict: A dictionary mapping requirement IDs to user stories.
        """
        user_stories = {}
        
        if self.ai_model:
            try:
                for req in self.requirements.get('requirements', []):
                    prompt = PromptTemplate(
                        input_variables=["requirement"],
                        template="Generate a user story based on this requirement:\n{requirement}\n\nUser Story:"
                    )
                    
                    # Convert requirement to string for the prompt
                    req_str = json.dumps(req, indent=2)
                    
                    # Generate user story using AI
                    user_stories[req['id']] = self.ai_model(prompt.format(requirement=req_str))
            except Exception as e:
                self.logger.error(f"Error generating user stories with AI: {str(e)}")
                # Fall back to template-based generation
                user_stories = self._generate_template_user_stories()
        else:
            # Use template-based generation
            user_stories = self._generate_template_user_stories()
        
        return user_stories
    
    def _generate_template_user_stories(self):
        """
        Generate template-based user stories for each requirement.
        
        Returns:
            dict: A dictionary mapping requirement IDs to user stories.
        """
        user_stories = {}
        
        for req in self.requirements.get('requirements', []):
            req_id = req['id']
            req_type = req['type'].replace('_', ' ').title()
            req_desc = req['description']
            req_details = req['details']
            
            if req_type == "Data Extraction":
                user_story = f"""
                As a business user,
                I want to automatically extract data from {req_desc.lower().split('extract')[1] if 'extract' in req_desc.lower() else 'the source'},
                So that I can {req_details.lower().split('should')[1] if 'should' in req_details.lower() else 'process it efficiently'}.
                
                Acceptance Criteria:
                - The bot should extract all required data fields
                - The extraction should be accurate and reliable
                - The bot should handle exceptions gracefully
                """
            elif req_type == "Data Processing":
                user_story = f"""
                As a business user,
                I want to automatically process {req_desc.lower().split('process')[1] if 'process' in req_desc.lower() else 'the data'},
                So that I can {req_details.lower().split('should')[1] if 'should' in req_details.lower() else 'use it for business purposes'}.
                
                Acceptance Criteria:
                - The bot should apply all business rules correctly
                - The processing should be efficient and reliable
                - The bot should handle exceptions gracefully
                """
            elif req_type == "System Integration":
                user_story = f"""
                As a business user,
                I want to integrate with {req_desc.lower().split('with')[1] if 'with' in req_desc.lower() else 'the target system'},
                So that I can {req_details.lower().split('should')[1] if 'should' in req_details.lower() else 'ensure data consistency across systems'}.
                
                Acceptance Criteria:
                - The bot should connect to the system securely
                - The data exchange should be accurate and reliable
                - The bot should handle connection issues gracefully
                """
            elif req_type == "Process Automation":
                user_story = f"""
                As a business user,
                I want to automate {req_desc.lower().split('automate')[1] if 'automate' in req_desc.lower() else 'the process'},
                So that I can {req_details.lower().split('should')[1] if 'should' in req_details.lower() else 'improve efficiency and reduce manual effort'}.
                
                Acceptance Criteria:
                - The bot should execute the process end-to-end
                - The automation should be reliable and consistent
                - The bot should provide appropriate logging and notifications
                """
            else:
                user_story = f"""
                As a business user,
                I want to {req_desc.lower()},
                So that I can {req_details.lower().split('should')[1] if 'should' in req_details.lower() else 'improve business operations'}.
                
                Acceptance Criteria:
                - The bot should fulfill the requirement accurately
                - The implementation should be reliable and efficient
                - The bot should handle exceptions appropriately
                """
            
            user_stories[req_id] = user_story.strip()
        
        return user_stories